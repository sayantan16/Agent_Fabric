"""
File Content Reader
Reads actual file contents instead of just metadata
This is CRITICAL for the orchestrator to see real data
"""

import os
import json
import pandas as pd
import PyPDF2
from typing import Dict, List, Any, Optional
import chardet
import openpyxl
import docx
import yaml


class FileContentReader:
    """Reads actual file contents for the orchestrator to see real data."""

    def __init__(self):
        self.max_preview_rows = 100  # For large files
        self.max_text_preview = 5000  # Characters for text files

    def read_file_contents(
        self, file_path: str, file_type: str = None
    ) -> Dict[str, Any]:
        """
        Read actual file contents based on file type.

        Returns:
            Dict with actual data, not just metadata
        """

        result = {
            "path": file_path,
            "type": file_type or self._detect_file_type(file_path),
            "read_success": False,
            "content": None,
            "structure": None,
            "error": None,
            "size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
        }

        try:
            # CSV Files
            if file_path.endswith(".csv") or file_type == "text/csv":
                result.update(self._read_csv(file_path))

            # Excel Files
            elif (
                file_path.endswith((".xlsx", ".xls"))
                or "spreadsheet" in str(file_type).lower()
            ):
                result.update(self._read_excel(file_path))

            # JSON Files
            elif file_path.endswith(".json") or file_type == "application/json":
                result.update(self._read_json(file_path))

            # Text Files
            elif file_path.endswith(".txt") or file_type == "text/plain":
                result.update(self._read_text(file_path))

            # PDF Files
            elif file_path.endswith(".pdf") or file_type == "application/pdf":
                result.update(self._read_pdf(file_path))

            # Word Documents
            elif file_path.endswith((".docx", ".doc")):
                result.update(self._read_word(file_path))

            # YAML Files
            elif file_path.endswith((".yml", ".yaml")):
                result.update(self._read_yaml(file_path))

            # Python Files
            elif file_path.endswith(".py"):
                result.update(self._read_code(file_path, "python"))

            # Default: Try as text
            else:
                result.update(self._read_text(file_path))

        except Exception as e:
            result["error"] = str(e)
            result["read_success"] = False

        return result

    def _read_csv(self, file_path: str) -> Dict:
        """Read CSV with pandas and extract meaningful data."""
        try:
            # Read CSV with intelligent parsing
            df = pd.read_csv(file_path, nrows=self.max_preview_rows)

            # Get full dataframe for stats
            df_full = pd.read_csv(file_path)

            return {
                "read_success": True,
                "structure": "tabular",
                "content": {
                    "columns": df.columns.tolist(),
                    "total_rows": len(df_full),
                    "total_columns": len(df.columns),
                    "first_10_rows": df.head(10).to_dict("records"),
                    "last_5_rows": df.tail(5).to_dict("records"),
                    "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                    "sample_values": {
                        col: df[col].dropna().head(5).tolist() for col in df.columns
                    },
                    "null_counts": df_full.isnull().sum().to_dict(),
                    "numeric_columns": df.select_dtypes(
                        include=["number"]
                    ).columns.tolist(),
                    "text_columns": df.select_dtypes(
                        include=["object"]
                    ).columns.tolist(),
                    "statistics": df.describe().to_dict() if not df.empty else {},
                },
            }
        except Exception as e:
            return {"error": f"CSV read error: {str(e)}", "read_success": False}

    def _read_excel(self, file_path: str) -> Dict:
        """Read Excel files with multiple sheets support."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}

            for sheet_name in excel_file.sheet_names[:5]:  # Limit to first 5 sheets
                df = pd.read_excel(
                    file_path, sheet_name=sheet_name, nrows=self.max_preview_rows
                )
                sheets_data[sheet_name] = {
                    "columns": df.columns.tolist(),
                    "rows": len(df),
                    "preview": df.head(10).to_dict("records"),
                    "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                }

            return {
                "read_success": True,
                "structure": "tabular",
                "content": {
                    "sheet_count": len(excel_file.sheet_names),
                    "sheet_names": excel_file.sheet_names,
                    "sheets": sheets_data,
                    "primary_sheet": (
                        sheets_data[excel_file.sheet_names[0]]
                        if excel_file.sheet_names
                        else {}
                    ),
                },
            }
        except Exception as e:
            return {"error": f"Excel read error: {str(e)}", "read_success": False}

    def _read_json(self, file_path: str) -> Dict:
        """Read JSON files and understand structure."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            # Analyze JSON structure
            structure_info = self._analyze_json_structure(data)

            return {
                "read_success": True,
                "structure": "json",
                "content": {
                    "data": data,
                    "type": type(data).__name__,
                    "structure_info": structure_info,
                    "keys": list(data.keys()) if isinstance(data, dict) else None,
                    "length": len(data) if isinstance(data, (list, dict)) else None,
                    "preview": str(data)[:1000] if len(str(data)) > 1000 else data,
                },
            }
        except Exception as e:
            return {"error": f"JSON read error: {str(e)}", "read_success": False}

    def _read_text(self, file_path: str) -> Dict:
        """Read text files with encoding detection."""
        try:
            # Detect encoding
            with open(file_path, "rb") as f:
                raw_data = f.read(10000)
                detected = chardet.detect(raw_data)
                encoding = detected["encoding"] or "utf-8"

            # Read with detected encoding
            with open(file_path, "r", encoding=encoding, errors="ignore") as f:
                text = f.read()

            # Analyze text
            lines = text.split("\n")
            words = text.split()

            return {
                "read_success": True,
                "structure": "text",
                "content": {
                    "text": text[: self.max_text_preview],
                    "full_text": text,
                    "full_length": len(text),
                    "line_count": len(lines),
                    "word_count": len(words),
                    "encoding": encoding,
                    "first_lines": lines[:20],
                    "has_headers": self._detect_headers(lines),
                    "appears_structured": self._detect_structure(lines),
                },
            }
        except Exception as e:
            return {"error": f"Text read error: {str(e)}", "read_success": False}

    def _read_pdf(self, file_path: str) -> Dict:
        """Read PDF files and extract text."""
        try:
            text_content = []
            metadata = {}

            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)

                # Extract text from first few pages
                for i in range(min(5, num_pages)):
                    page = pdf_reader.pages[i]
                    text_content.append(page.extract_text())

                # Get metadata
                if pdf_reader.metadata:
                    metadata = {
                        "title": pdf_reader.metadata.get("/Title", ""),
                        "author": pdf_reader.metadata.get("/Author", ""),
                        "subject": pdf_reader.metadata.get("/Subject", ""),
                    }

            full_text = "\n".join(text_content)

            return {
                "read_success": True,
                "structure": "pdf",
                "content": {
                    "text_preview": full_text[: self.max_text_preview],
                    "page_count": num_pages,
                    "metadata": metadata,
                    "first_pages_text": text_content,
                    "word_count": len(full_text.split()),
                    "has_text": bool(full_text.strip()),
                },
            }
        except Exception as e:
            return {"error": f"PDF read error: {str(e)}", "read_success": False}

    def _read_word(self, file_path: str) -> Dict:
        """Read Word documents."""
        try:
            doc = docx.Document(file_path)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract tables
            tables = []
            for table in doc.tables[:5]:  # First 5 tables
                table_data = []
                for row in table.rows[:10]:  # First 10 rows
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            full_text = "\n".join(paragraphs)

            return {
                "read_success": True,
                "structure": "document",
                "content": {
                    "text": full_text[: self.max_text_preview],
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                    "first_paragraphs": paragraphs[:10],
                    "tables_preview": tables,
                    "word_count": len(full_text.split()),
                },
            }
        except Exception as e:
            return {"error": f"Word read error: {str(e)}", "read_success": False}

    def _read_yaml(self, file_path: str) -> Dict:
        """Read YAML configuration files."""
        try:
            with open(file_path, "r") as f:
                data = yaml.safe_load(f)

            return {
                "read_success": True,
                "structure": "yaml",
                "content": {
                    "data": data,
                    "type": type(data).__name__,
                    "keys": list(data.keys()) if isinstance(data, dict) else None,
                    "preview": str(data)[:1000],
                },
            }
        except Exception as e:
            return {"error": f"YAML read error: {str(e)}", "read_success": False}

    def _read_code(self, file_path: str, language: str) -> Dict:
        """Read code files with syntax awareness."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                code = f.read()

            lines = code.split("\n")

            # Basic code analysis
            import_lines = [
                l for l in lines if l.strip().startswith(("import ", "from "))
            ]
            function_lines = [l for l in lines if "def " in l or "class " in l]

            return {
                "read_success": True,
                "structure": "code",
                "content": {
                    "language": language,
                    "code": code[: self.max_text_preview],
                    "full_code": code,
                    "line_count": len(lines),
                    "imports": import_lines,
                    "functions_classes": function_lines,
                    "has_main": "__main__" in code,
                },
            }
        except Exception as e:
            return {"error": f"Code read error: {str(e)}", "read_success": False}

    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension."""
        ext = os.path.splitext(file_path)[1].lower()
        type_map = {
            ".csv": "text/csv",
            ".json": "application/json",
            ".txt": "text/plain",
            ".pdf": "application/pdf",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xls": "application/vnd.ms-excel",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".py": "text/x-python",
            ".yml": "text/yaml",
            ".yaml": "text/yaml",
        }
        return type_map.get(ext, "application/octet-stream")

    def _analyze_json_structure(
        self, data: Any, depth: int = 0, max_depth: int = 3
    ) -> Dict:
        """Analyze JSON structure recursively."""
        if depth >= max_depth:
            return {"type": type(data).__name__, "truncated": True}

        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys())[:10],
                "key_count": len(data),
                "sample": {
                    k: self._analyze_json_structure(v, depth + 1)
                    for k, v in list(data.items())[:3]
                },
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "item_type": (
                    self._analyze_json_structure(data[0], depth + 1) if data else None
                ),
            }
        else:
            return {"type": type(data).__name__, "value": str(data)[:100]}

    def _detect_headers(self, lines: List[str]) -> bool:
        """Detect if text file has headers."""
        if not lines:
            return False

        first_line = lines[0]
        # Common header patterns
        return any(
            [
                "," in first_line and len(first_line.split(",")) > 2,
                "\t" in first_line and len(first_line.split("\t")) > 2,
                "|" in first_line and len(first_line.split("|")) > 2,
                first_line.startswith("#"),
                first_line.isupper(),
            ]
        )

    def _detect_structure(self, lines: List[str]) -> bool:
        """Detect if text appears to be structured data."""
        if len(lines) < 3:
            return False

        # Check for consistent delimiters
        delimiters = [",", "\t", "|", ";"]
        for delim in delimiters:
            counts = [line.count(delim) for line in lines[:10] if line.strip()]
            if counts and all(c == counts[0] and c > 0 for c in counts):
                return True

        return False

    def process_all_files(self, files: List[Dict]) -> List[Dict]:
        """
        Process all uploaded files and read their contents.

        Args:
            files: List of file metadata dicts from upload

        Returns:
            List of files with actual content included
        """

        enriched_files = []

        for file_info in files:
            file_path = file_info.get("path", "")
            file_type = file_info.get("type", "")

            if os.path.exists(file_path):
                # Read actual content
                content_data = self.read_file_contents(file_path, file_type)

                # Merge with original metadata
                enriched_file = {**file_info, **content_data}
                enriched_files.append(enriched_file)

                print(
                    f"✓ Read {file_info.get('original_name', 'file')}: "
                    f"{content_data.get('structure', 'unknown')} structure, "
                    f"{'success' if content_data.get('read_success') else 'failed'}"
                )
            else:
                # File doesn't exist
                enriched_file = {
                    **file_info,
                    "read_success": False,
                    "error": f"File not found: {file_path}",
                }
                enriched_files.append(enriched_file)
                print(f"✗ File not found: {file_path}")

        return enriched_files
